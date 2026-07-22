"""Text extraction for uploaded plant documents."""

from __future__ import annotations

import io
from pathlib import Path


def parse_bytes(filename: str, data: bytes) -> tuple[str, int]:
    """Return (text, page_count_estimate)."""
    name = (filename or "upload.txt").lower()
    suffix = Path(name).suffix

    if suffix == ".pdf":
        return _parse_pdf(data)
    if suffix in {".docx"}:
        return _parse_docx(data)
    if suffix in {".xlsx", ".xls"}:
        return _parse_xlsx(data)
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        # Image OCR not wired; store filename hint so upload still lands in KB
        return (
            f"[Image upload: {filename}]\nPlant image attached. Ask AI with this file context later.",
            1,
        )
    if suffix in {".eml"}:
        return data.decode("utf-8", errors="ignore"), 1

    # .txt and unknown text-like
    try:
        return data.decode("utf-8"), 1
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore"), 1


def _parse_pdf(data: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    return text or "[PDF had no extractable text]", max(1, len(reader.pages))


def _parse_docx(data: bytes) -> tuple[str, int]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text or "[DOCX had no extractable text]", max(1, len(doc.paragraphs) // 40)


def _parse_xlsx(data: bytes) -> tuple[str, int]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines) or "[Spreadsheet empty]", max(1, len(wb.worksheets))
